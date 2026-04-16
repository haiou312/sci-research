#!/usr/bin/env python3
"""
Generate a branded SPD Bank docx from briefing-curator Agent output.

Uses a template docx to preserve headers (SPD Bank logo), footers
(geometric decoration + page numbers), styles, and numbering definitions.
Clears the body content and injects new content matching the template structure.

Usage:
  generate-branded-docx.py \
    --template path/to/briefing-template.docx \
    --input path/to/curator-output.txt \
    --output path/to/output.docx

Exit codes:
  0 — success
  1 — missing python-docx
  2 — template not found
  3 — input parse error
  4 — docx generation error
"""

import argparse
import re
import sys
from copy import deepcopy
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip3 install python-docx", file=sys.stderr)
    sys.exit(1)


# ---------------------------------------------------------------------------
# 1. Parse curator output
# ---------------------------------------------------------------------------

def parse_curator_output(text):
    """Parse the structured output from briefing-curator agent."""
    result = {"title": "", "date": "", "toc": [], "stories": [], "references": [], "disclaimer": ""}

    # Extract DATE
    m = re.search(r"^DATE:\s*(.+)$", text, re.M)
    if m:
        result["date"] = m.group(1).strip()

    # Extract TOC
    toc_match = re.search(r"^TOC:\s*\n((?:- .+\n?)+)", text, re.M)
    if toc_match:
        result["toc"] = [line.lstrip("- ").strip() for line in toc_match.group(1).strip().split("\n") if line.strip()]

    # Extract STORIES
    stories_match = re.search(r"^STORIES:\s*\n(.+?)(?=^REFERENCES:)", text, re.M | re.S)
    if stories_match:
        stories_text = stories_match.group(1).strip()
        # Split by **N. pattern
        story_blocks = re.split(r"\*\*(\d+)\.", stories_text)
        # story_blocks[0] is empty/preamble, then alternating: number, rest
        i = 1
        while i < len(story_blocks) - 1:
            num = story_blocks[i].strip()
            rest = story_blocks[i + 1].strip()
            # rest starts with "title**\n\nbody [N]"
            title_match = re.match(r"(.+?)\*\*\s*\n+(.+)", rest, re.S)
            if title_match:
                title = title_match.group(1).strip()
                body = title_match.group(2).strip()
                result["stories"].append({"num": int(num), "title": title, "body": body})
            i += 2

    # Extract REFERENCES
    refs_match = re.search(r"^REFERENCES:\s*\n(.+?)(?=\nDISCLAIMER:)", text, re.M | re.S)
    if refs_match:
        for line in refs_match.group(1).strip().split("\n"):
            line = line.strip()
            if line:
                # [N]  URL
                ref_m = re.match(r"\[(\d+)\]\s+(.+)", line)
                if ref_m:
                    result["references"].append({"num": int(ref_m.group(1)), "url": ref_m.group(2).strip()})

    # Extract DISCLAIMER
    disc_match = re.search(r"^DISCLAIMER:\s*\n(.+)", text, re.M | re.S)
    if disc_match:
        result["disclaimer"] = disc_match.group(1).strip()

    if not result["stories"]:
        print("ERROR: no stories parsed from input", file=sys.stderr)
        sys.exit(3)

    return result


# ---------------------------------------------------------------------------
# 2. DocX generation helpers
# ---------------------------------------------------------------------------

def set_run_font(run, font_name="FangSong_GB2312", size_pt=None, bold=False):
    """Set font properties on a run, including East Asian font."""
    run.bold = bold
    if size_pt:
        run.font.size = Pt(size_pt)
    run.font.name = font_name
    # Set East Asian font via XML
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def add_numbering_to_paragraph(paragraph, num_id=8, ilvl=0):
    """Add numbered/bulleted list properties to a paragraph via XML."""
    pPr = paragraph._element.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl_elem = OxmlElement("w:ilvl")
    ilvl_elem.set(qn("w:val"), str(ilvl))
    numId_elem = OxmlElement("w:numId")
    numId_elem.set(qn("w:val"), str(num_id))
    numPr.append(ilvl_elem)
    numPr.append(numId_elem)
    pPr.append(numPr)


# ---------------------------------------------------------------------------
# 3. Core generation
# ---------------------------------------------------------------------------

def generate_docx(template_path, data, output_path):
    """Generate branded docx from parsed data using template."""
    try:
        doc = Document(str(template_path))
    except Exception as e:
        print(f"ERROR: cannot open template: {e}", file=sys.stderr)
        sys.exit(2)

    body = doc.element.body

    # --- Save section properties (CRITICAL for header/footer preservation) ---
    # Body-level sectPr
    body_sectPrs = []
    for sp in body.findall(qn("w:sectPr")):
        body_sectPrs.append(deepcopy(sp))

    # Paragraph-embedded sectPr (section breaks within document)
    para_sectPrs = []
    for p in body.findall(qn("w:p")):
        pPr = p.find(qn("w:pPr"))
        if pPr is not None:
            sp = pPr.find(qn("w:sectPr"))
            if sp is not None:
                para_sectPrs.append(deepcopy(sp))

    # --- Clear all body content ---
    for child in list(body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("p", "tbl", "sdt"):
            body.remove(child)

    # --- Add content ---

    # 1. Title line: "新闻简报  时间：2026年4月16日"
    p_title = doc.add_paragraph()
    run_label = p_title.add_run("新闻简报")
    set_run_font(run_label, bold=True, size_pt=16)
    run_space = p_title.add_run(" 时间：")
    set_run_font(run_space, size_pt=14)
    run_date = p_title.add_run(data["date"])
    set_run_font(run_date, size_pt=14)

    # 2. Empty paragraph separator
    doc.add_paragraph()

    # 3. TOC (bullet list)
    for toc_item in data["toc"]:
        p_toc = doc.add_paragraph(style="ListParagraph")
        add_numbering_to_paragraph(p_toc, num_id=8, ilvl=0)
        run = p_toc.add_run(toc_item)
        set_run_font(run, bold=True, size_pt=10.5)

    # 4. Empty paragraph after TOC
    doc.add_paragraph()

    # 5. Story bodies
    for story in data["stories"]:
        # Empty line before story
        doc.add_paragraph()

        # Title paragraph: "**N.title**"
        p_story_title = doc.add_paragraph()
        run_title = p_story_title.add_run(f"{story['num']}.{story['title']}")
        set_run_font(run_title, bold=True, size_pt=12)

        # Body paragraph
        p_body = doc.add_paragraph()
        run_body = p_body.add_run(story["body"])
        set_run_font(run_body, size_pt=10.5)

    # 6. Insert section break (paragraph-embedded sectPr) before references
    if para_sectPrs:
        last_p = body.findall(qn("w:p"))[-1]
        last_pPr = last_p.find(qn("w:pPr"))
        if last_pPr is None:
            last_pPr = OxmlElement("w:pPr")
            last_p.insert(0, last_pPr)
        last_pPr.append(para_sectPrs[0])

    # 7. References heading
    doc.add_paragraph()
    p_ref_heading = doc.add_paragraph()
    run_ref = p_ref_heading.add_run("References")
    set_run_font(run_ref, bold=True, size_pt=12)

    # 8. References table (borderless, matching template)
    if data["references"]:
        table = doc.add_table(rows=len(data["references"]), cols=2)
        # Remove all borders for a clean look
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        tblBorders = OxmlElement("w:tblBorders")
        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "none")
            border.set(qn("w:sz"), "0")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "auto")
            tblBorders.append(border)
        # Remove existing borders if any
        existing = tblPr.find(qn("w:tblBorders"))
        if existing is not None:
            tblPr.remove(existing)
        tblPr.append(tblBorders)

        for i, ref in enumerate(data["references"]):
            row = table.rows[i]
            # Column 1: [N]
            cell_num = row.cells[0]
            cell_num.text = ""
            p_num = cell_num.paragraphs[0]
            run_num = p_num.add_run(f"[{ref['num']}]")
            set_run_font(run_num, size_pt=10)
            # Set column width narrow
            cell_num.width = Pt(40)

            # Column 2: URL
            cell_url = row.cells[1]
            cell_url.text = ""
            p_url = cell_url.paragraphs[0]
            run_url = p_url.add_run(ref["url"])
            set_run_font(run_url, size_pt=10)

    # 9. Empty paragraphs before disclaimer
    doc.add_paragraph()
    doc.add_paragraph()

    # 10. Disclaimer
    if data["disclaimer"]:
        p_disc = doc.add_paragraph()
        # Bold "免责声明：" prefix
        disc_text = data["disclaimer"]
        if disc_text.startswith("免责声明：") or disc_text.startswith("免责声明:"):
            prefix = "免责声明：" if "：" in disc_text[:6] else "免责声明:"
            rest = disc_text[len(prefix):]
            run_prefix = p_disc.add_run(prefix)
            set_run_font(run_prefix, bold=True, size_pt=12)
            run_rest = p_disc.add_run(rest)
            set_run_font(run_rest, size_pt=12)
        else:
            run_disc = p_disc.add_run(disc_text)
            set_run_font(run_disc, size_pt=12)

    # --- Restore body-level sectPr ---
    for sp in body_sectPrs:
        body.append(sp)

    # --- Save ---
    try:
        doc.save(str(output_path))
    except Exception as e:
        print(f"ERROR: failed to save docx: {e}", file=sys.stderr)
        sys.exit(4)


# ---------------------------------------------------------------------------
# 4. Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Generate branded SPD Bank docx from curator output.")
    parser.add_argument("--template", required=True, help="Path to briefing-template.docx")
    parser.add_argument("--input", required=True, help="Path to curator output text file")
    parser.add_argument("--output", required=True, help="Output docx path")
    args = parser.parse_args()

    template_path = Path(args.template).expanduser()
    if not template_path.exists():
        print(f"ERROR: template not found: {template_path}", file=sys.stderr)
        sys.exit(2)

    input_path = Path(args.input).expanduser()
    if not input_path.exists():
        print(f"ERROR: input file not found: {input_path}", file=sys.stderr)
        sys.exit(3)

    text = input_path.read_text(encoding="utf-8")
    data = parse_curator_output(text)

    output_path = Path(args.output).expanduser()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    generate_docx(template_path, data, output_path)
    print(f"OK: generated {output_path} with {len(data['stories'])} stories")


if __name__ == "__main__":
    main()
