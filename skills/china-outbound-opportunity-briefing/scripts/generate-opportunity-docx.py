#!/usr/bin/env python3
"""Generate an institutional DOCX from opportunity-briefing Markdown."""

from __future__ import annotations

import argparse
from io import BytesIO
import re
from pathlib import Path
import sys
from typing import Any
from urllib.error import HTTPError, URLError
from urllib.parse import urlparse
from urllib.request import Request, urlopen

try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    requirements = Path(__file__).resolve().parents[3] / "requirements.txt"
    print(
        "ERROR: python-docx is not installed. Install or update it with:\n"
        f"  python3 -m pip install --user --upgrade -r {requirements}",
        file=sys.stderr,
    )
    raise SystemExit(1)


PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
BODY_FONT = "Arial Unicode MS"
EAST_ASIA_FONT = BODY_FONT
INK = RGBColor(11, 37, 69)
BLUE = RGBColor(0, 91, 172)
RED = RGBColor(200, 16, 46)
GRAY = RGBColor(90, 98, 108)
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
LIGHT_RED = "FBE9ED"
WHITE = "FFFFFF"


def set_run_font(
    run,
    *,
    size: float = 10.5,
    bold: bool | None = None,
    italic: bool | None = None,
    color: RGBColor | None = None,
    ascii_font: str = BODY_FONT,
    east_asia_font: str = EAST_ASIA_FONT,
) -> None:
    run.font.name = ascii_font
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia_font)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = "D7DCE2", size: int = 4) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != PAGE_WIDTH_DXA:
        raise ValueError(f"table widths must sum to {PAGE_WIDTH_DXA}: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(PAGE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run = OxmlElement("w:r")
    field_run.append(begin)
    field_run.append(instruction)
    field_run.append(separate)
    field_run.append(value)
    field_run.append(end)
    paragraph._p.append(field_run)
    suffix = paragraph.add_run(" 页")
    set_run_font(suffix, size=9, color=GRAY)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, INK, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = BODY_FONT
        style.font.size = Pt(10.5)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [6000, 3360])
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
    left = table.cell(0, 0).paragraphs[0]
    left.paragraph_format.space_after = Pt(0)
    run = left.add_run("SPD BANK  |  中资企业商机拓展简报")
    set_run_font(run, size=8.5, bold=True, color=BLUE)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.space_after = Pt(0)
    run = right.add_run("内部使用")
    set_run_font(run, size=8.5, bold=True, color=RED)

    footer = section.footer
    add_page_field(footer.paragraphs[0])


def add_hyperlink(paragraph, text: str, url: str, *, size: float = 10) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "005BAC")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), BODY_FONT)
    rfonts.set(qn("w:hAnsi"), BODY_FONT)
    rfonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    size_node = OxmlElement("w:sz")
    size_node.set(qn("w:val"), str(int(size * 2)))
    rpr.extend([rfonts, color, underline, size_node])
    run.append(rpr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|\[[^\]]+\]\([^)]+\))")


def add_inline_markdown(paragraph, text: str, *, size: float = 10.5, color=None) -> None:
    position = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        else:
            link = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link:
                add_hyperlink(paragraph, link.group(1), link.group(2), size=size)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=size, color=color)


def strip_markdown(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    return text.strip()


def parse_table(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) < 2:
        return rows[0] if rows else [], []
    separator = rows[1]
    if all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in separator):
        return rows[0], rows[2:]
    return rows[0], rows[1:]


def table_widths(headers: list[str]) -> list[int]:
    key = tuple(headers)
    if key == ("优先级", "企业/项目", "最新事件", "潜在需求", "建议动作", "时间窗口"):
        return [850, 1450, 1850, 1650, 2150, 1410]
    if key == (
        "英国实体",
        "Company No.",
        "中国母公司/关联方",
        "本期变化",
        "中资置信度",
        "商业意义",
    ):
        return [1500, 1100, 1600, 2000, 900, 2260]
    if key == ("日期/窗口", "事项", "涉及企业/行业", "关注原因", "建议跟进"):
        return [1100, 1800, 1800, 2400, 2260]
    count = max(1, len(headers))
    base = PAGE_WIDTH_DXA // count
    widths = [base] * count
    widths[-1] += PAGE_WIDTH_DXA - sum(widths)
    return widths


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    if not headers:
        return
    table = doc.add_table(rows=1, cols=len(headers))
    widths = table_widths(headers)
    set_table_geometry(table, widths)
    set_table_borders(table)
    mark_header_row(table.rows[0])

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_GRAY)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(strip_markdown(header))
        set_run_font(run, size=9, bold=True, color=INK)

    for row_values in rows:
        row = table.add_row()
        values = row_values + [""] * (len(headers) - len(row_values))
        for index, value in enumerate(values[: len(headers)]):
            cell = row.cells[index]
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            if index == 0 or headers[index] in {
                "Company No.",
                "中资置信度",
                "日期/窗口",
                "时间窗口",
            }:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(strip_markdown(value))
            set_run_font(run, size=8.5, color=INK)
            if index == 0 and value.strip() == "高":
                set_cell_shading(cell, LIGHT_RED)
                run.bold = True
                run.font.color.rgb = RED
            elif index == 0 and value.strip() in {"中", "观察"}:
                set_cell_shading(cell, LIGHT_BLUE)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def load_image(source: str, max_bytes: int = 8 * 1024 * 1024) -> BytesIO:
    parsed = urlparse(source)
    if parsed.scheme in {"http", "https"}:
        request = Request(
            source,
            headers={"User-Agent": "sci-research-opportunity-docx/1"},
        )
        with urlopen(request, timeout=12) as response:
            content_length = response.headers.get("Content-Length")
            if content_length and int(content_length) > max_bytes:
                raise ValueError("image exceeds size limit")
            payload = response.read(max_bytes + 1)
        if len(payload) > max_bytes:
            raise ValueError("image exceeds size limit")
        return BytesIO(payload)
    path = Path(source).expanduser()
    if not path.exists():
        raise FileNotFoundError(source)
    payload = path.read_bytes()
    if len(payload) > max_bytes:
        raise ValueError("image exceeds size limit")
    return BytesIO(payload)


def add_image(doc: Document, caption: str, source: str) -> bool:
    try:
        image = load_image(source)
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run().add_picture(image, width=Inches(6.1))
        caption_paragraph = doc.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_paragraph.paragraph_format.space_before = Pt(0)
        caption_paragraph.paragraph_format.space_after = Pt(4)
        run = caption_paragraph.add_run(caption)
        set_run_font(run, size=9, italic=True, color=GRAY)
        return True
    except (OSError, ValueError, HTTPError, URLError, UnidentifiedImageError):
        return False
    except Exception:
        return False


try:
    from docx.image.exceptions import UnrecognizedImageError as UnidentifiedImageError
except ImportError:
    UnidentifiedImageError = ValueError


def add_title_block(doc: Document, title: str, metadata: list[str]) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(14)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(5)
    run = kicker.add_run("CHINA OUTBOUND  |  UK & EUROPE")
    set_run_font(run, size=9.5, bold=True, color=RED)

    title_paragraph = doc.add_paragraph()
    title_paragraph.paragraph_format.space_before = Pt(0)
    title_paragraph.paragraph_format.space_after = Pt(7)
    title_paragraph.paragraph_format.keep_with_next = True
    run = title_paragraph.add_run(title)
    set_run_font(run, size=24, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("英国经济、欧洲出海、跨境并购、投资布局与英国实体变化")
    set_run_font(run, size=12, color=BLUE)

    if metadata:
        meta = doc.add_paragraph()
        meta.paragraph_format.space_before = Pt(0)
        meta.paragraph_format.space_after = Pt(10)
        for index, value in enumerate(metadata):
            if index:
                run = meta.add_run("  |  ")
                set_run_font(run, size=9.5, color=GRAY)
            add_inline_markdown(meta, value, size=9.5, color=GRAY)


def parse_front_matter(lines: list[str]) -> tuple[str, list[str], int]:
    title = "中资企业商机拓展简报"
    metadata = []
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if line.startswith("# "):
            title = strip_markdown(line[2:])
            index += 1
            continue
        if line.startswith("> "):
            metadata.append(strip_markdown(line[2:]))
            index += 1
            continue
        if not line:
            index += 1
            continue
        break
    return title, metadata, index


def render_markdown(doc: Document, lines: list[str], start_index: int) -> tuple[int, int]:
    index = start_index
    image_count = 0
    image_failures = 0
    paragraph_buffer: list[str] = []
    in_disclaimer_section = False

    def flush_paragraph() -> None:
        if not paragraph_buffer:
            return
        text = " ".join(item.strip() for item in paragraph_buffer).strip()
        paragraph_buffer.clear()
        if text:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.widow_control = True
            if in_disclaimer_section:
                paragraph.paragraph_format.space_after = Pt(4)
                paragraph.paragraph_format.line_spacing = 1.05
                add_inline_markdown(paragraph, text, size=9.5)
            else:
                add_inline_markdown(paragraph, text)

    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            flush_paragraph()
            index += 1
            continue

        if line.startswith("## "):
            flush_paragraph()
            heading = strip_markdown(line[3:])
            in_disclaimer_section = heading == "来源与免责声明"
            doc.add_heading(heading, level=1)
            index += 1
            continue

        if line.startswith("### "):
            flush_paragraph()
            doc.add_heading(strip_markdown(line[4:]), level=2)
            index += 1
            continue

        image_match = re.fullmatch(r"!\[([^\]]*)\]\(([^)]+)\)", line)
        if image_match:
            flush_paragraph()
            if add_image(doc, image_match.group(1), image_match.group(2)):
                image_count += 1
            else:
                image_failures += 1
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(
                    f"配图未嵌入：{image_match.group(1)} | {image_match.group(2)}"
                )
                set_run_font(run, size=9, italic=True, color=GRAY)
            index += 1
            continue

        if line.startswith("|"):
            flush_paragraph()
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            headers, rows = parse_table(table_lines)
            add_table(doc, headers, rows)
            continue

        if line.startswith("- "):
            flush_paragraph()
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(paragraph, line[2:])
            index += 1
            continue

        if line.startswith("*图片来源：") and line.endswith("*"):
            flush_paragraph()
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(6)
            add_inline_markdown(paragraph, line.strip("*"), size=8.5, color=GRAY)
            index += 1
            continue

        if line.startswith("> "):
            flush_paragraph()
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.2)
            paragraph.paragraph_format.right_indent = Inches(0.2)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(6)
            add_inline_markdown(paragraph, line[2:], color=GRAY)
            index += 1
            continue

        paragraph_buffer.append(line)
        index += 1

    flush_paragraph()
    return image_count, image_failures


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    input_path = args.input.expanduser()
    output_path = args.output.expanduser()
    if not input_path.is_file():
        print(f"ERROR: input not found: {input_path}", file=sys.stderr)
        return 2

    text = input_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    title, metadata, start_index = parse_front_matter(lines)

    doc = Document()
    configure_document(doc)
    add_title_block(doc, title, metadata)
    image_count, image_failures = render_markdown(doc, lines, start_index)

    core_properties = doc.core_properties
    core_properties.title = title
    core_properties.subject = "UK and Europe commercial opportunity intelligence"
    core_properties.author = "SPD Bank Sci-Research"
    core_properties.keywords = (
        "China outbound, United Kingdom, Europe, M&A, Companies House"
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(output_path)
    except Exception as exc:
        print(f"ERROR: failed to save DOCX: {exc}", file=sys.stderr)
        return 4
    print(
        f"OK: generated {output_path} images={image_count} "
        f"image_fallbacks={image_failures}"
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
